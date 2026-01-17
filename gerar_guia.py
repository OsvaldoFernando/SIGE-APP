from docx import Document
from docx.shared import Inches
import datetime

def gerar_guia_sige():
    document = Document()
    document.add_heading('Guia da Aplicação - SIGE (Sistema Integrado de Gestão Educacional)', 0)

    document.add_heading('1. Visão Geral', level=1)
    document.add_paragraph(
        'O SIGE é um sistema robusto desenvolvido em Django focado na gestão educacional moderna, '
        'abrangendo desde a admissão de candidatos até o lançamento de notas e gestão administrativa.'
    )

    document.add_heading('2. Módulo de Admissão e Matrícula', level=1)
    document.add_paragraph(
        'Localizado em /nova-matricula/, este módulo permite gerir todo o fluxo de entrada de estudantes.'
    )
    document.add_heading('2.1 Notas de Admissão', level=2)
    document.add_paragraph(
        'Interface centralizada para lançamento de notas de testes de admissão com dashboard de analytics em tempo real, '
        'incluindo médias de turma e taxa de aptidão.'
    )

    document.add_heading('2.2 URLs Amigáveis (ERP-Style)', level=2)
    document.add_paragraph(
        'As URLs de inscrição foram modernizadas para usar slugs em vez de IDs numéricos. '
        'Exemplo: /inscricao/engenharia-informatica/ em vez de /inscricao/2/.'
    )

    document.add_heading('3. Gestão de Empates (Configurações)', level=1)
    document.add_paragraph(
        'O sistema possui uma gestão de empates altamente configurável nas Configurações Académicas Globais.'
    )
    document.add_paragraph('Critérios disponíveis:', style='List Bullet')
    document.add_paragraph('Mais Velho primeiro', style='List Bullet')
    document.add_paragraph('Mais Novo primeiro', style='List Bullet')
    document.add_paragraph('Ordem de Inscrição', style='List Bullet')
    document.add_paragraph('Prioridade Feminino/Masculino', style='List Bullet')
    document.add_paragraph('Combinados: Feminino e Mais Novo / Feminino e Mais Velho', style='List Bullet')

    document.add_heading('4. Gestão Administrativa', level=1)
    document.add_paragraph('Controle de cursos, disciplinas, turmas e recursos humanos.')

    document.add_heading('5. Atualizações Recentes', level=1)
    document.add_paragraph(f'Documento atualizado em: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}')

    document.save('Guia_Aplicacao_SIGE.docx')
    print("Guia gerado com sucesso!")

if __name__ == "__main__":
    gerar_guia_sige()
